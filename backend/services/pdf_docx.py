import os
import re
import base64
import urllib.request
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, Table, TableStyle, PageBreak, NextPageTemplate, FrameBreak, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_SECTION

import zlib

def fetch_mermaid_image(mermaid_code: str, filepath: str) -> bool:
    try:
        encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
        url = f"https://mermaid.ink/img/{encoded}?bgColor=FFFFFF"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(filepath, 'wb') as out_file:
            out_file.write(response.read())
        return True
    except Exception as e:
        print(f"Error fetching diagram: {e}")
        return False

def fetch_mermaid_png(mermaid_code: str, filepath: str) -> bool:
    try:
        compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
        url = f"https://kroki.io/mermaid/png/{encoded}"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req) as response, open(filepath, 'wb') as out_file:
            out_file.write(response.read())
        return True
    except Exception as e:
        print(f"Error fetching diagram PNG: {e}")
        return False

def parse_markdown_blocks(content_text):
    blocks = []
    lines = content_text.strip().split('\n')
    current_para = []
    in_table = False
    table_rows = []
    
    for line in lines:
        line_s = line.strip()
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', line_s)
        if img_match:
            if current_para:
                blocks.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            blocks.append({'type': 'image', 'caption': img_match.group(1), 'url': img_match.group(2)})
            continue
            
        if line_s.startswith('|') and line_s.endswith('|'):
            if current_para:
                blocks.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
            in_table = True
            cells = [c.strip() for c in line_s.split('|')[1:-1]]
            if not all(re.match(r'^[-:]+$', c) for c in cells):
                table_rows.append(cells)
            continue
        else:
            if in_table:
                blocks.append({'type': 'table', 'rows': table_rows})
                table_rows = []
                in_table = False
                
        if line_s == "":
            if current_para:
                blocks.append({'type': 'paragraph', 'text': ' '.join(current_para)})
                current_para = []
        else:
            current_para.append(line_s)
            
    if current_para:
        blocks.append({'type': 'paragraph', 'text': ' '.join(current_para)})
    if in_table:
        blocks.append({'type': 'table', 'rows': table_rows})
    return blocks

class DocumentCompilerService:
    @staticmethod
    def compile_pdf(paper_data: dict, filepath: str) -> str:
        diagrams_injected_local = False
        base_dir = os.path.dirname(filepath)
        os.makedirs(base_dir, exist_ok=True)
        
        doc = BaseDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36
        )
        
        frame_width = (doc.width - 18) / 2
        
        # Increase top_height for IEEE Access: Title, Authors, Abstract, Index Terms
        top_height = 4.5 * inch
        body_height = doc.height - top_height
        
        is_custom = paper_data.get("format") == "Custom Advanced"
        
        frame_top = Frame(doc.leftMargin, doc.bottomMargin + body_height, doc.width, top_height, id='top')
        
        story = []
        if is_custom:
            frame_body_first = Frame(doc.leftMargin, doc.bottomMargin, doc.width, body_height, id='body_first')
            frame_full = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id='full')
            doc.addPageTemplates([
                PageTemplate(id='FirstPage', frames=[frame_top, frame_body_first]),
                PageTemplate(id='Body', frames=[frame_full])
            ])
            story.append(NextPageTemplate('Body'))
            img_width = doc.width
        else:
            frame_col1_first = Frame(doc.leftMargin, doc.bottomMargin, frame_width, body_height, id='col1_first')
            frame_col2_first = Frame(doc.leftMargin + frame_width + 18, doc.bottomMargin, frame_width, body_height, id='col2_first')
            frame_col1 = Frame(doc.leftMargin, doc.bottomMargin, frame_width, doc.height, id='col1')
            frame_col2 = Frame(doc.leftMargin + frame_width + 18, doc.bottomMargin, frame_width, doc.height, id='col2')
            doc.addPageTemplates([
                PageTemplate(id='FirstPage', frames=[frame_top, frame_col1_first, frame_col2_first]),
                PageTemplate(id='Body', frames=[frame_col1, frame_col2])
            ])
            story.append(NextPageTemplate('Body'))
            img_width = frame_width
        
        ieee_blue = colors.HexColor("#0071A5")
        accent_color = colors.HexColor("#000000")
        
        title_style = ParagraphStyle(name='AcademicTitle', fontName='Helvetica-Bold', fontSize=24, leading=26, textColor=ieee_blue, alignment=1, spaceAfter=18)
        authors_style = ParagraphStyle(name='AcademicAuthors', fontName='Helvetica-Bold', fontSize=10, leading=12, textColor=colors.black, alignment=1, spaceAfter=2)
        inst_style = ParagraphStyle(name='AcademicInst', fontName='Times-Roman', fontSize=9, leading=11, textColor=colors.black, alignment=1, spaceAfter=18)
        
        abstract_body_style = ParagraphStyle(name='AbstractBody', fontName='Times-Roman', fontSize=9, leading=12, textColor=colors.black, alignment=4, spaceAfter=18)
        
        heading1_style = ParagraphStyle(name='AcademicH1', fontName='Helvetica-Bold', fontSize=10, leading=12, textColor=ieee_blue, alignment=1, spaceBefore=12, spaceAfter=6, keepWithNext=True)
        body_style = ParagraphStyle(name='AcademicBody', fontName='Times-Roman', fontSize=10, leading=12, textColor=colors.black, firstLineIndent=12, spaceAfter=0, alignment=4)
        caption_style = ParagraphStyle(name='CaptionStyle', fontName='Times-Bold', fontSize=8, leading=10, textColor=colors.black, alignment=1, spaceBefore=4, spaceAfter=12)

        # 1. Document Title
        story.append(Paragraph("<b>IEEE Access</b>", ParagraphStyle('Logo', fontName='Helvetica-Bold', fontSize=28, textColor=ieee_blue, alignment=0, spaceAfter=24)))
        story.append(Paragraph(paper_data.get("title", "Untitled Research Paper"), title_style))
        
        # 2. Authors (IEEE Side-by-Side Layout)
        author_raw = paper_data.get("authorName", "Sabarishwaran")
        if not author_raw or author_raw == "Academic Researcher":
            author_raw = "Sabarishwaran"
        inst_raw = paper_data.get("institution", "AI Research Institute")
        author_email = paper_data.get("authorEmail", "")
        
        authors_list = [a.strip() for a in author_raw.split(',')]
        
        # Build cells for each author
        author_cells = []
        for author_name in authors_list:
            cell_story = []
            cell_story.append(Paragraph(author_name, authors_style))
            cell_story.append(Paragraph(inst_raw, inst_style))
            if author_email:
                cell_story.append(Paragraph(f"<u><font color='blue'>{author_email}</font></u>", inst_style))
            author_cells.append(cell_story)
            
        # Group cells into rows of max_cols
        max_cols = min(3, len(author_cells)) if len(author_cells) > 0 else 1
        author_rows = []
        for i in range(0, len(author_cells), max_cols):
            row = author_cells[i:i+max_cols]
            # Pad the row with empty strings if less than max_cols
            while len(row) < max_cols:
                row.append("")
            author_rows.append(row)
            
        if author_rows:
            author_table = Table(author_rows, colWidths=[doc.width / max_cols] * max_cols)
            author_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 12)
            ]))
            story.append(author_table)
            story.append(Spacer(1, 12))
        
        # 3. Abstract and Keywords
        abstract_text = paper_data.get("abstract", "No abstract provided.")
        story.append(Paragraph(f"<b><i>Abstract</i>—</b> {abstract_text}", abstract_body_style))
        
        keywords = paper_data.get("keywords", "Satellite Image Change Detection, Deep Learning, Multi-Temporal Imagery, Artificial Intelligence, Autonomous Systems.")
        story.append(Paragraph(f"<b>Keywords:</b> {keywords}", abstract_body_style))
        
        # Jump into 2 columns
        story.append(FrameBreak())
        
        # Sections
        sections = paper_data.get("sections", {})
        sequence = [(k, k.upper()) for k in sections.keys()]
            
        for idx, (key, display_title) in enumerate(sequence):
            content = sections.get(key) or sections.get(key.lower()) or sections.get(key.replace(" ", ""))
            if content:
                story.append(Paragraph(display_title, heading1_style))
                blocks = parse_markdown_blocks(content)
                for block in blocks:
                    if block['type'] == 'paragraph':
                        story.append(Paragraph(block['text'], body_style))
                    elif block['type'] == 'image':
                        try:
                            from reportlab.lib.utils import ImageReader
                            ir = ImageReader(block['url'])
                            iw, ih = ir.getSize()
                            img_height = img_width * (ih / float(iw)) if iw > 0 else img_width * 0.75
                            img = RLImage(block['url'], width=img_width, height=img_height)
                            img.hAlign = 'CENTER'
                            story.append(Spacer(1, 6))
                            story.append(img)
                            story.append(Paragraph(block['caption'], caption_style))
                        except Exception:
                            story.append(Paragraph(f"[Image Placeholder: {block['caption']}]", caption_style))
                    elif block['type'] == 'table':
                        story.append(Spacer(1, 6))
                        table_data = [[Paragraph(cell, ParagraphStyle('TC', fontName='Times-Roman', fontSize=9, alignment=1)) for cell in row] for row in block['rows']]
                        if table_data:
                            t = Table(table_data, colWidths=[img_width/len(table_data[0])]*len(table_data[0]))
                            t.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,0), 1, ieee_blue), ('LINEBELOW', (0,0), (-1,0), 1, ieee_blue), ('LINEBELOW', (0,-1), (-1,-1), 1, ieee_blue), ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke), ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0), (-1,-1), 'MIDDLE')]))
                            story.append(t)
                            story.append(Spacer(1, 6))
                
                inject_now = False
                if is_custom:
                    if "structure" in key.lower() or "architecture" in key.lower() or "agentic ai" in key.lower() or "system" in key.lower() or "design" in key.lower():
                        inject_now = True
                else:
                    if "methodology" in key.lower() or "architecture" in key.lower() or "results" in key.lower():
                        inject_now = True
                
                if inject_now and not diagrams_injected_local:
                    diagrams = paper_data.get("diagrams", [])
                    if diagrams:
                        diagrams_injected_local = True
                        for d_idx, d in enumerate(diagrams):
                            d_code = d.get('mermaidCode', '')
                            if d_code:
                                img_path = os.path.join(base_dir, f"diag_{idx}_{d_idx}.png")
                                if fetch_mermaid_png(d_code, img_path):
                                    story.append(Spacer(1, 12))
                                    try:
                                        from reportlab.platypus import Image
                                        from reportlab.lib.utils import ImageReader
                                        img_obj = ImageReader(img_path)
                                        iw, ih = img_obj.getSize()
                                        aspect = ih / float(iw)
                                        final_width = 400
                                        final_height = final_width * aspect
                                        img_flowable = Image(img_path, width=final_width, height=final_height)
                                        img_flowable.hAlign = 'CENTER'
                                        story.append(img_flowable)
                                    except Exception as e:
                                        print(f"PNG render failed: {e}")
                                    cap = d.get('caption', f"Figure {d_idx + 1}")
                                    story.append(Spacer(1, 6))
                                    story.append(Paragraph(f"{cap}: {d.get('title', '')}", caption_style))

        diagrams = paper_data.get("diagrams", [])
        if diagrams and not diagrams_injected_local:
            diagrams_injected_local = True
            for d_idx, d in enumerate(diagrams):
                d_code = d.get('mermaidCode', '')
                if d_code:
                    img_path = os.path.join(base_dir, f"diag_fallback_{d_idx}.png")
                    if fetch_mermaid_png(d_code, img_path):
                        story.append(Spacer(1, 12))
                        try:
                            from reportlab.platypus import Image
                            from reportlab.lib.utils import ImageReader
                            img_obj = ImageReader(img_path)
                            iw, ih = img_obj.getSize()
                            aspect = ih / float(iw)
                            final_width = 400
                            final_height = final_width * aspect
                            img_flowable = Image(img_path, width=final_width, height=final_height)
                            img_flowable.hAlign = 'CENTER'
                            story.append(img_flowable)
                        except Exception as e:
                            print(f"PNG render failed: {e}")
                        cap = d.get('caption', f"Figure {d_idx + 1}")
                        story.append(Spacer(1, 6))
                        story.append(Paragraph(f"{cap}: {d.get('title', '')}", caption_style))

        story.append(Paragraph("REFERENCES", heading1_style))
        refs = paper_data.get("references", [])
        if not refs:
            refs = ["[1] A. Vaswani et al., 'Attention is all you need,' NIPS, 2017."]
        
        ref_style = ParagraphStyle(name='AcademicRef', fontName='Times-Roman', fontSize=8, leading=10, textColor=colors.black, leftIndent=15, firstLineIndent=-15, spaceAfter=4, alignment=4)
        for ref in refs:
            story.append(Paragraph(ref, ref_style))
            
        doc.build(story)
        return filepath

    @staticmethod
    def compile_docx(paper_data: dict, filepath: str) -> str:
        base_dir = os.path.dirname(filepath)
        os.makedirs(base_dir, exist_ok=True)
        doc = Document()
        
        diagrams_injected_docx_local = False
        
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '1')
        
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(paper_data.get("title", "Untitled Research Paper"))
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0x00, 0x71, 0xA5)
        title_p.alignment = 1 
        
        author_raw = paper_data.get("authorName", "Sabarishwaran")
        if not author_raw or author_raw == "Academic Researcher":
            author_raw = "Sabarishwaran"
        inst_raw = paper_data.get("institution", "AI Research Institute")
        authors_list = [a.strip().upper() for a in author_raw.split(',')]
            
        authors_str = ", ".join(authors_list)
        p = doc.add_paragraph()
        r1 = p.add_run(authors_str + '\n')
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(10)
        
        inst_text = inst_raw
        author_email = paper_data.get("authorEmail", "")
        if author_email:
            inst_text += f"\n{author_email}"
            
        r2 = p.add_run(inst_text)
        r2.italic = True
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(9)
        p.alignment = 1 
            
        doc.add_paragraph()
        
        abs_p = doc.add_paragraph()
        abs_run1 = abs_p.add_run("Abstract— ")
        abs_run1.font.name = 'Times New Roman'
        abs_run1.font.size = Pt(10)
        abs_run1.bold = True
        abs_run1.italic = True
        abs_run1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        
        abs_run2 = abs_p.add_run(paper_data.get('abstract', ''))
        abs_run2.font.name = 'Times New Roman'
        abs_run2.font.size = Pt(9)
        abs_p.alignment = 3 
        
        kw_p = doc.add_paragraph()
        kw_run1 = kw_p.add_run("Keywords: ")
        kw_run1.font.name = 'Times New Roman'
        kw_run1.font.size = Pt(9)
        kw_run1.bold = True
        
        keywords = paper_data.get("keywords", "Satellite Image Change Detection, Deep Learning, Multi-Temporal Imagery, Artificial Intelligence, Autonomous Systems.")
        kw_run2 = kw_p.add_run(keywords)
        kw_run2.font.name = 'Times New Roman'
        kw_run2.font.size = Pt(9)
        kw_p.alignment = 3
        
        doc.add_paragraph() 
        
        is_custom = paper_data.get("format") == "Custom Advanced"
        
        if not is_custom:
            new_section = doc.add_section(WD_SECTION.CONTINUOUS)
            new_sectPr = new_section._sectPr
            new_cols = new_sectPr.xpath('./w:cols')[0]
            new_cols.set(qn('w:num'), '2')
            new_cols.set(qn('w:space'), '360')
        
        sections = paper_data.get("sections", {})
        sequence = [(k, k.upper()) for k in sections.keys()]
            
        for idx, (key, display_title) in enumerate(sequence):
            content = sections.get(key) or sections.get(key.lower()) or sections.get(key.replace(" ", ""))
            if content:
                h = doc.add_paragraph()
                hr = h.add_run(display_title)
                hr.font.name = 'Arial'
                hr.font.size = Pt(10)
                hr.font.color.rgb = RGBColor(0x00, 0x71, 0xA5)
                h.alignment = 1 
                
                blocks = parse_markdown_blocks(content)
                for block in blocks:
                    if block['type'] == 'paragraph':
                        p = doc.add_paragraph()
                        pr = p.add_run(block['text'])
                        pr.font.name = 'Times New Roman'
                        pr.font.size = Pt(10)
                        p.alignment = 3
                    elif block['type'] == 'image':
                        try:
                            p_img = doc.add_paragraph()
                            p_img.alignment = 1
                            r_img = p_img.add_run()
                            r_img.add_picture(block['url'], width=Inches(3.0))
                            p = doc.add_paragraph()
                            pr = p.add_run(block['caption'])
                            pr.bold = True
                            pr.font.name = 'Times New Roman'
                            pr.font.size = Pt(8)
                            p.alignment = 1
                        except Exception:
                            p = doc.add_paragraph()
                            p.add_run(f"[Image Placeholder: {block['caption']}]").font.size = Pt(8)
                            p.alignment = 1
                    elif block['type'] == 'table':
                        if block['rows']:
                            table = doc.add_table(rows=len(block['rows']), cols=len(block['rows'][0]))
                            table.style = 'Light Shading'
                            for row_idx, row_data in enumerate(block['rows']):
                                for col_idx, cell_text in enumerate(row_data):
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = cell_text
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.font.name = 'Times New Roman'
                                            r.font.size = Pt(9)
                
                inject_now = False
                if is_custom:
                    if "structure" in key.lower() or "architecture" in key.lower() or "agentic ai" in key.lower() or "system" in key.lower() or "design" in key.lower():
                        inject_now = True
                else:
                    if "methodology" in key.lower() or "architecture" in key.lower() or "results" in key.lower():
                        inject_now = True
                        
                if inject_now and not diagrams_injected_docx_local:
                    diagrams = paper_data.get("diagrams", [])
                    if diagrams:
                        diagrams_injected_docx_local = True
                        for d_idx, d in enumerate(diagrams):
                            d_code = d.get('mermaidCode')
                            if d_code:
                                img_path = os.path.join(base_dir, f"diag_docx_{idx}_{d_idx}.png")
                                if fetch_mermaid_png(d_code, img_path):
                                    p_img = doc.add_paragraph()
                                    p_img.alignment = 1
                                    r_img = p_img.add_run()
                                    r_img.add_picture(img_path, width=Inches(6.0 if is_custom else 3.0))
                                    p = doc.add_paragraph()
                                    cap = d.get('caption', f"Figure {d_idx + 1}")
                                    pr = p.add_run(f"{cap}: {d.get('title', '')}")
                                    pr.bold = True
                                    pr.font.name = 'Times New Roman'
                                    pr.font.size = Pt(8)
                                    p.alignment = 1
                                    
        # Ensure diagrams are injected even if no section name matched the keywords for DOCX
        diagrams = paper_data.get("diagrams", [])
        if diagrams and not diagrams_injected_docx_local:
            diagrams_injected_docx_local = True
            for d_idx, d in enumerate(diagrams):
                d_code = d.get('mermaidCode')
                if d_code:
                    img_path = os.path.join(base_dir, f"diag_docx_fallback_{d_idx}.png")
                    if fetch_mermaid_png(d_code, img_path):
                        p_img = doc.add_paragraph()
                        p_img.alignment = 1
                        r_img = p_img.add_run()
                        r_img.add_picture(img_path, width=Inches(6.0 if is_custom else 3.0))
                        p = doc.add_paragraph()
                        cap = d.get('caption', f"Figure {d_idx + 1}")
                        pr = p.add_run(f"{cap}: {d.get('title', '')}")
                        pr.bold = True
                        pr.font.name = 'Times New Roman'
                        pr.font.size = Pt(8)
                        p.alignment = 1

        h = doc.add_paragraph()
        hr = h.add_run("REFERENCES")
        hr.font.name = 'Arial'
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0x00, 0x71, 0xA5)
        h.alignment = 1 # Center aligned
        
        refs = paper_data.get("references", [])
        for ref in refs:
            p = doc.add_paragraph()
            pr = p.add_run(ref)
            pr.font.name = 'Times New Roman'
            pr.font.size = Pt(8)
            
        doc.save(filepath)
        if hasattr(DocumentCompilerService, '_diagrams_injected_docx'):
            delattr(DocumentCompilerService, '_diagrams_injected_docx')
        return filepath
